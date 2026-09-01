"""DOCX processing using python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from processing.base import BaseProcessor
from processing.exceptions import DocumentProcessingError
from schemas.document import (
    ContentBlock,
    ContentBlockType,
    DocumentStatus,
    FileType,
    NormalizedDocument,
)


class DOCXProcessor(BaseProcessor):
    def process(
        self,
        file_path: Path,
        document_id: str,
        filename: str,
    ) -> NormalizedDocument:
        try:
            document = Document(str(file_path))
        except PackageNotFoundError as exc:
            raise DocumentProcessingError(f"Could not open DOCX: {exc}") from exc
        except Exception as exc:  # noqa: BLE001
            raise DocumentProcessingError(f"Could not open DOCX: {exc}") from exc

        content: list[ContentBlock] = []

        try:
            location = 0
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                location += 1
                content.append(
                    ContentBlock(
                        type=ContentBlockType.PARAGRAPH,
                        location=location,
                        text=text,
                    )
                )

            for table_index, table in enumerate(document.tables):
                table_text = self._table_to_text(table)
                if not table_text:
                    continue
                location += 1
                content.append(
                    ContentBlock(
                        type=ContentBlockType.TABLE,
                        location=location,
                        text=table_text,
                        extra={"table_index": table_index},
                    )
                )
        except Exception as exc:  # noqa: BLE001
            raise DocumentProcessingError(
                f"Failed to extract DOCX content: {exc}"
            ) from exc

        paragraph_count = sum(
            1 for block in content if block.type == ContentBlockType.PARAGRAPH
        )
        table_count = sum(
            1 for block in content if block.type == ContentBlockType.TABLE
        )

        return NormalizedDocument(
            document_id=document_id,
            filename=filename,
            file_type=FileType.DOCX,
            status=DocumentStatus.PROCESSED,
            content=content,
            metadata={
                "paragraph_count": paragraph_count,
                "table_count": table_count,
            },
        )

    @staticmethod
    def _table_to_text(table) -> str:
        """Convert a docx table into a normalized, pipe-delimited text block."""
        rows_text = []
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            if any(cells_text):
                rows_text.append(" | ".join(cells_text))
        return "\n".join(rows_text)
